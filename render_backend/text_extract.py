"""
Shared upload text-extraction helpers. No ML dependencies here on purpose —
both api/predict and api/predict_narrow import this, so keeping it light
means it never contributes to either function's bundle size beyond
pdfplumber / python-docx themselves.
"""

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}


def extract_text_from_pdf(file_stream) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_stream) -> str:
    import docx

    document = docx.Document(file_stream)
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text_from_txt(file_stream) -> str:
    raw = file_stream.read()
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="ignore")
    return raw


def extract_text(filename: str, file_stream) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_stream)
    if ext == "docx":
        return extract_text_from_docx(file_stream)
    if ext == "txt":
        return extract_text_from_txt(file_stream)
    raise ValueError(f"Unsupported file type: .{ext}")
