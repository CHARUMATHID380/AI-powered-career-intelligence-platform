"""
CareerCast — Resume -> Suitable Job Web App
============================================
Flask backend that wraps the Logistic Regression resume classifier
(resume_job_classifier.joblib) in a small web app. Users paste resume text
or upload a .pdf / .docx / .txt file, and the app returns a ranked list of
job categories with match percentages.

RUN LOCALLY:
    1. pip install -r requirements.txt
    2. Make sure resume_job_classifier.joblib sits next to this file
       (produced by train_model.py — see that script's docstring).
    3. python app.py
    4. Open http://localhost:5000
"""

import io
import os
import re
import string

from flask import Flask, jsonify, render_template, request

import joblib
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# ---------------------------------------------------------------------------
# Setup — identical text-cleaning logic used at training time. This MUST stay
# in sync with train_model.py, otherwise predictions will be inaccurate.
# ---------------------------------------------------------------------------
nltk.download("stopwords", quiet=True)
nltk.download("wordnet", quiet=True)
nltk.download("omw-1.4", quiet=True)

STOPWORDS = set(stopwords.words("english"))
LEMMATIZER = WordNetLemmatizer()

URL_RE = re.compile(r"http\S+|www\.\S+")
EMAIL_RE = re.compile(r"\S+@\S+")
PHONE_RE = re.compile(r"\+?\d[\d\-\s()]{7,}\d")
NON_ALPHA_RE = re.compile(r"[^a-z\s]")
MULTI_SPACE_RE = re.compile(r"\s+")

MODEL_PATH = os.path.join(BASE_DIR, "resume_job_classifier.joblib")
TOP_N = 5

SKILL_KEYWORDS = [
    "Python", "Java", "C++", "C#", "JavaScript", "TypeScript", "SQL", "R",
    "HTML", "CSS", "PHP", "Go", "Ruby", "Swift", "Kotlin", "Scala",
    "Machine Learning", "Deep Learning", "Data Analysis", "Data Visualization",
    "Data Science", "Natural Language Processing", "Computer Vision",
    "Artificial Intelligence", "Statistics", "TensorFlow", "PyTorch",
    "Scikit-learn", "Pandas", "NumPy", "Keras",
    "React", "Angular", "Vue", "Node.js", "Django", "Flask", "Spring Boot",
    "REST API", "GraphQL", "Microservices",
    "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes", "DevOps",
    "CI/CD", "Git", "Linux",
    "Excel", "Power BI", "Tableau", "MySQL", "PostgreSQL", "MongoDB",
    "Project Management", "Agile", "Scrum", "Communication", "Leadership",
    "Teamwork", "Problem Solving", "Time Management",
    "Accounting", "Finance", "Marketing", "Sales", "SEO", "Content Writing",
    "Business Analysis", "Financial Modeling", "Auditing",
]
_SKILL_PATTERNS = [
    (skill, re.compile(r"(?<![a-zA-Z])" + re.escape(skill).replace(r"\ ", r"\s+") + r"(?![a-zA-Z])", re.IGNORECASE))
    for skill in SKILL_KEYWORDS
]


def extract_skills(raw_text: str, limit: int = 14):
    found = []
    for skill, pattern in _SKILL_PATTERNS:
        if pattern.search(raw_text):
            found.append(skill)
    return found[:limit]

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_CONTENT_LENGTH = 8 * 1024 * 1024  # 8 MB

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH

_model = None
_model_error = None
try:
    _model = joblib.load(MODEL_PATH)
except Exception as exc:  # noqa: BLE001
    _model_error = str(exc)


def clean_resume_text(text: str) -> str:
    text = text.lower()
    text = URL_RE.sub(" ", text)
    text = EMAIL_RE.sub(" ", text)
    text = PHONE_RE.sub(" ", text)
    text = text.translate(str.maketrans("", "", string.punctuation))
    text = NON_ALPHA_RE.sub(" ", text)
    tokens = text.split()
    tokens = [
        LEMMATIZER.lemmatize(tok)
        for tok in tokens
        if tok not in STOPWORDS and len(tok) > 2
    ]
    text = " ".join(tokens)
    text = MULTI_SPACE_RE.sub(" ", text).strip()
    return text


def predict_top_jobs(resume_text: str, n: int = TOP_N):
    cleaned = clean_resume_text(resume_text)
    if not cleaned:
        return []
    proba = _model.predict_proba([cleaned])[0]
    classes = _model.classes_
    ranked = sorted(zip(classes, proba), key=lambda x: x[1], reverse=True)[:n]
    return [{"job": job, "score": round(float(score) * 100, 2)} for job, score in ranked]


# ---------------------------------------------------------------------------
# Text extraction for uploaded files
# ---------------------------------------------------------------------------
def extract_text_from_pdf(file_stream) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_stream) -> str:
    import docx

    document = docx.Document(file_stream)
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text_from_txt(file_stream) -> str:
    raw = file_stream.read()
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="ignore")
    return raw


def extract_text(filename: str, file_stream) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_stream)
    if ext == "docx":
        return extract_text_from_docx(file_stream)
    if ext == "txt":
        return extract_text_from_txt(file_stream)
    raise ValueError(f"Unsupported file type: .{ext}")


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html", model_ready=_model is not None, model_error=_model_error)


@app.route("/api/predict", methods=["POST"])
def predict():
    if _model is None:
        return jsonify({
            "error": "Model not loaded. Place resume_job_classifier.joblib next to app.py and restart the server."
        }), 500

    resume_text = ""

    # Case 1: a file was uploaded
    if "resume_file" in request.files and request.files["resume_file"].filename:
        f = request.files["resume_file"]
        filename = f.filename
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in ALLOWED_EXTENSIONS:
            return jsonify({"error": f"Unsupported file type .{ext}. Use PDF, DOCX, or TXT."}), 400
        try:
            file_bytes = f.read()
            resume_text = extract_text(filename, io.BytesIO(file_bytes))
        except Exception as exc:  # noqa: BLE001
            return jsonify({"error": f"Could not read file: {exc}"}), 400

    # Case 2: pasted text
    elif request.form.get("resume_text"):
        resume_text = request.form.get("resume_text", "")

    else:
        return jsonify({"error": "Paste resume text or upload a PDF/DOCX/TXT file."}), 400

    if not resume_text or not resume_text.strip():
        return jsonify({"error": "No readable text was found in the resume."}), 400

    try:
        results = predict_top_jobs(resume_text, n=TOP_N)
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": f"Prediction failed: {exc}"}), 500

    if not results:
        return jsonify({"error": "Resume text was too short or empty after cleaning."}), 400

    preview = resume_text.strip().replace("\r", "")
    if len(preview) > 600:
        preview = preview[:600] + "…"

    skills = extract_skills(resume_text)

    return jsonify({"results": results, "preview": preview, "skills": skills})


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
